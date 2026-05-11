from __future__ import annotations

import re
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "LAB_MANUAL.md"
TARGET = ROOT / "LAB_MANUAL.docx"

DOCUMENT_TITLE = "AI in Cyber Security Laboratory Manual"
DARK_BLUE = RGBColor(18, 59, 99)
BLACK = RGBColor(0, 0, 0)
BODY = RGBColor(31, 31, 31)
MUTED = RGBColor(95, 95, 95)
CODE_BACKGROUND = "F0F0F0"
TERMINAL_BACKGROUND = "F7F7F7"


def clean_text(text: str) -> str:
    replacements = {
        "Ã¢Å“â€¦": "[OK]",
        "Ã¢Å¡Â Ã¯Â¸Â": "[Warning]",
        "Ã¢Å¡Â ": "[Warning]",
        "Ã°Å¸â€Â¥": "[High]",
        "âœ…": "[OK]",
        "âš ï¸": "[Warning]",
        "âš ": "[Warning]",
        "ðŸ”¥": "[High]",
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return text.strip()


def markdown_to_plain(text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", r"\1", text)
    text = re.sub(r"\[`?([^`\]]+)`?\]\(([^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return clean_text(text.replace("`", ""))


def set_cell_background(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    tc_pr.append(shading)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_document(document: Document):
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    header = section.header.paragraphs[0]
    header.text = DOCUMENT_TITLE
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.name = "Helvetica"
    header.runs[0].font.size = Pt(9)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"{DOCUMENT_TITLE} | Page ")
    add_page_number(footer)
    for run in footer.runs:
        run.font.name = "Helvetica"
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED


def configure_styles(document: Document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Helvetica"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY

    for name in ("Program Title", "Subheading", "File Path", "Code Text"):
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    program = styles["Program Title"]
    program.font.name = "Helvetica"
    program.font.bold = True
    program.font.size = Pt(18)
    program.font.color.rgb = DARK_BLUE
    program.paragraph_format.space_before = Pt(14)
    program.paragraph_format.space_after = Pt(12)
    program.paragraph_format.keep_with_next = True

    subheading = styles["Subheading"]
    subheading.font.name = "Helvetica"
    subheading.font.bold = True
    subheading.font.size = Pt(14)
    subheading.font.color.rgb = DARK_BLUE
    subheading.paragraph_format.space_before = Pt(12)
    subheading.paragraph_format.space_after = Pt(7)
    subheading.paragraph_format.keep_with_next = True

    path = styles["File Path"]
    path.font.name = "Helvetica"
    path.font.italic = True
    path.font.size = Pt(9)
    path.font.color.rgb = MUTED
    path.paragraph_format.space_after = Pt(8)

    code = styles["Code Text"]
    code.font.name = "Courier New"
    code.font.size = Pt(8)
    code.paragraph_format.space_after = Pt(0)


def add_heading(document: Document, text: str):
    document.add_paragraph(markdown_to_plain(text), style="Program Title")


def add_subheading(document: Document, text: str):
    label = "Terminal Output" if text.strip().lower() == "output" else text
    document.add_paragraph(markdown_to_plain(label), style="Subheading")


def add_body_text(document: Document, text: str):
    if not text.strip():
        return
    style = "File Path" if text.startswith(("File:", "Terminal output file:")) else None
    paragraph = document.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if style is None else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.add_run(markdown_to_plain(text))


def add_bullet(document: Document, text: str):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.add_run(markdown_to_plain(text[2:]))


def wrap_monospace(text: str, max_chars: int = 88) -> str:
    wrapped: list[str] = []
    for raw_line in text.rstrip("\n").splitlines():
        line = raw_line.rstrip()
        if not line:
            wrapped.append("")
            continue
        indent = re.match(r"^\s*", line).group(0)
        wrapped.extend(
            textwrap.wrap(
                line,
                width=max_chars,
                subsequent_indent=indent + "    ",
                replace_whitespace=False,
                drop_whitespace=False,
                break_long_words=True,
                break_on_hyphens=False,
                tabsize=4,
            )
            or [""]
        )
    return "\n".join(wrapped)


def add_monospace_block(document: Document, text: str, background: str):
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(6.27)
    cell = table.cell(0, 0)
    set_cell_background(cell, background)
    set_cell_margins(cell)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.style = document.styles["Code Text"]
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.add_run(wrap_monospace(text))
    document.add_paragraph()


def add_code_block(document: Document, code: str):
    add_monospace_block(document, code, CODE_BACKGROUND)


def add_terminal_output(document: Document, output: str):
    add_monospace_block(document, output, TERMINAL_BACKGROUND)


def add_image(document: Document, image_path: Path, caption: str):
    if not image_path.exists():
        add_body_text(document, f"Image not found: {image_path}")
        return

    caption_para = document.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.paragraph_format.keep_with_next = True
    caption_run = caption_para.add_run(markdown_to_plain(caption))
    caption_run.bold = True
    caption_run.font.name = "Helvetica"
    caption_run.font.size = Pt(10)

    image_para = document.add_paragraph()
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.paragraph_format.space_after = Pt(12)
    image_run = image_para.add_run()
    image_run.add_picture(str(image_path), width=Inches(6.27))


def add_table_of_contents(document: Document):
    document.add_paragraph("Table of Contents", style="Subheading")
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "separate")
    fld_char_3 = OxmlElement("w:fldChar")
    fld_char_3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)
    run._r.append(fld_char_3)
    document.add_page_break()


def convert():
    document = Document()
    configure_document(document)
    configure_styles(document)

    in_code = False
    code_language = ""
    code_lines: list[str] = []
    current_section = ""
    first_program_seen = False

    for line in SOURCE.read_text(encoding="utf-8", errors="replace").splitlines():
        stripped = line.rstrip()

        if stripped.startswith("```"):
            if in_code:
                content = "\n".join(code_lines)
                if code_language.lower() in {"text", "terminal", "output"} or current_section == "Output":
                    add_terminal_output(document, content)
                else:
                    add_code_block(document, content)
                code_lines = []
                code_language = ""
                in_code = False
            else:
                in_code = True
                code_language = stripped.strip("`").strip()
            continue

        if in_code:
            code_lines.append(line)
            continue

        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            add_image(document, ROOT / image_match.group(2), image_match.group(1))
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = clean_text(heading_match.group(2))
            current_section = text
            if level == 1:
                title = document.add_paragraph()
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title.paragraph_format.space_after = Pt(18)
                run = title.add_run(markdown_to_plain(text))
                run.bold = True
                run.font.name = "Helvetica"
                run.font.size = Pt(22)
                run.font.color.rgb = DARK_BLUE
                add_table_of_contents(document)
            elif level == 2 and text.lower().startswith("program"):
                if first_program_seen:
                    document.add_page_break()
                first_program_seen = True
                add_heading(document, text)
            elif level in {2, 3}:
                add_subheading(document, text)
            else:
                add_body_text(document, text)
            continue

        if stripped.startswith("- "):
            add_bullet(document, stripped)
        else:
            add_body_text(document, stripped)

    document.save(TARGET)
    print(f"Created {TARGET}")


if __name__ == "__main__":
    convert()
